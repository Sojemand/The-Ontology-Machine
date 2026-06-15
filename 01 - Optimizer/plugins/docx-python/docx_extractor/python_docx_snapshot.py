from __future__ import annotations

from pathlib import Path

from .types import WordDocumentSnapshot, WordParagraphSnapshot, WordStageError, WordTableCellSnapshot


def ensure_python_docx() -> None:
    try:
        from docx import Document  # noqa: F401
    except ImportError as exc:
        raise WordStageError("adapter.runtime", str(exc)) from exc


def load_python_docx_snapshot(source: Path) -> WordDocumentSnapshot:
    try:
        from docx import Document
    except ImportError as exc:
        raise WordStageError("adapter.runtime", str(exc)) from exc

    try:
        document = Document(str(source))
        paragraphs = tuple(paragraph_snapshots(document))
        table_col_counts = tuple(max((len(row.cells) for row in table.rows), default=0) for table in document.tables)
        table_cells = tuple(table_cell_snapshots(document))
        image_count = sum(1 for rel in document.part.rels.values() if "image" in rel.reltype)
        return WordDocumentSnapshot(
            paragraphs=paragraphs,
            table_cells=table_cells,
            table_count=len(document.tables),
            table_col_counts=table_col_counts,
            has_images=image_count > 0,
            image_count=image_count,
            author=(document.core_properties.author or None),
            last_modified_by=(document.core_properties.last_modified_by or None),
            has_track_changes=has_track_changes(document),
        )
    except WordStageError:
        raise
    except Exception as exc:
        raise WordStageError("adapter.load", str(exc)) from exc


def paragraph_snapshots(document) -> list[WordParagraphSnapshot]:
    paragraphs: list[WordParagraphSnapshot] = []
    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = getattr(getattr(paragraph, "style", None), "name", "") or ""
        font_size = next(
            (
                run.font.size.pt
                for run in paragraph.runs
                if getattr(getattr(run, "font", None), "size", None) is not None
            ),
            None,
        )
        paragraphs.append(
            WordParagraphSnapshot(
                index=index,
                text=text,
                style_name=style_name,
                bold=any(run.bold for run in paragraph.runs if run.bold is not None),
                font_size=font_size,
            )
        )
    return paragraphs


def table_cell_snapshots(document) -> list[WordTableCellSnapshot]:
    cells: list[WordTableCellSnapshot] = []
    for table_index, table in enumerate(document.tables):
        for row_index, row in enumerate(table.rows):
            for col_index, cell in enumerate(row.cells):
                text = cell.text.strip()
                if text:
                    cells.append(WordTableCellSnapshot(table_index=table_index, row_index=row_index, col_index=col_index, text=text))
    return cells


def has_track_changes(document) -> bool:
    body_xml = document.element.body.xml
    return "<w:ins " in body_xml or "<w:del " in body_xml


__all__ = ["ensure_python_docx", "load_python_docx_snapshot"]
